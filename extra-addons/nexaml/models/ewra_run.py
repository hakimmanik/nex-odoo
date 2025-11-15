# -*- coding: utf-8 -*-
# Part of NexAML. See LICENSE file for full copyright and licensing details.

import logging
import re

from odoo import api, fields, models, _
from odoo.exceptions import ValidationError

_logger = logging.getLogger(__name__)


class EwraRun(models.Model):
    """Enterprise-Wide Risk Assessment Run."""
    _name = 'nexaml.ewra.run'
    _description = 'EWRA Run'
    _order = 'period_end desc, id desc'
    _inherit = ['mail.thread', 'mail.activity.mixin']

    name = fields.Char(
        string='EWRA Name',
        required=True,
        copy=False,
        readonly=True,
        default=lambda self: _('New'),
        help='Unique EWRA run identifier'
    )
    status = fields.Selection(
        [('draft', 'Draft'),
         ('in_progress', 'In Progress'),
         ('completed', 'Completed'),
         ('archived', 'Archived')],
        string='Status',
        required=True,
        default='draft',
        tracking=True,
        help='EWRA run status'
    )
    period_start = fields.Date(
        string='Period Start',
        required=True,
        tracking=True,
        help='Assessment period start date'
    )
    period_end = fields.Date(
        string='Period End',
        required=True,
        tracking=True,
        help='Assessment period end date'
    )

    # Relations
    pillar_ids = fields.One2many(
        'nexaml.ewra.pillar',
        'run_id',
        string='Risk Pillars',
        help='Risk assessment pillars'
    )
    settings_snapshot_id = fields.Many2one(
        'nexaml.ewra.settings.snapshot',
        string='Settings Snapshot',
        help='Settings at time of assessment'
    )
    narrative_id = fields.Many2one(
        'nexaml.ewra.narrative',
        string='Narrative',
        help='Foreword and conclusion'
    )

    # Overall Risk Scores
    overall_inherent_score = fields.Float(
        string='Overall Inherent Risk',
        compute='_compute_overall_scores',
        store=True,
        help='Weighted average inherent risk (1.0-3.0)'
    )
    overall_residual_score = fields.Float(
        string='Overall Residual Risk',
        compute='_compute_overall_scores',
        store=True,
        help='Weighted average residual risk (1.0-3.0)'
    )
    overall_risk_level = fields.Selection(
        [('low', 'Low'),
         ('medium', 'Medium'),
         ('high', 'High')],
        string='Overall Risk Level',
        compute='_compute_overall_scores',
        store=True,
        help='Overall risk classification'
    )

    # Compliance Officer Signature
    approved_by = fields.Many2one(
        'res.users',
        string='Approved By',
        help='Compliance officer who approved'
    )
    approved_date = fields.Datetime(
        string='Approved Date',
        help='When the EWRA was approved'
    )

    company_id = fields.Many2one(
        'res.company',
        string='Company',
        required=True,
        default=lambda self: self.env.company,
        help='Company'
    )

    @api.model_create_multi
    def create(self, vals_list):
        """Override create to generate EWRA run number."""
        for vals in vals_list:
            if vals.get('name', _('New')) == _('New'):
                vals['name'] = self.env['ir.sequence'].next_by_code('nexaml.ewra.run') or _('New')
        return super(EwraRun, self).create(vals_list)

    @api.depends('pillar_ids.inherent_score', 'pillar_ids.residual_score')
    def _compute_overall_scores(self):
        """Compute overall risk scores from pillars using weighted average.

        Pillar weights (matching nex-systems):
        - Customer: 40%
        - Geography: 20%
        - Products: 25%
        - Delivery: 15%
        - Supplier: 0% (excluded from overall calculation)
        """
        for run in self:
            if not run.pillar_ids:
                run.overall_inherent_score = 0.0
                run.overall_residual_score = 0.0
                run.overall_risk_level = 'low'
                continue

            # Define pillar weights (matching nex-systems exactly)
            pillar_weights = {
                'customer': 0.4,    # 40%
                'geography': 0.2,   # 20%
                'products': 0.25,   # 25%
                'delivery': 0.15,   # 15%
                'supplier': 0       # 0% - excluded from overall risk
            }

            # Filter out supplier pillar (weight = 0)
            core_pillars = run.pillar_ids.filtered(lambda p: p.pillar != 'supplier')

            if not core_pillars:
                run.overall_inherent_score = 0.0
                run.overall_residual_score = 0.0
                run.overall_risk_level = 'low'
                continue

            # Calculate total weight (only for pillars that exist)
            total_weight = sum([pillar_weights.get(p.pillar, 0) for p in core_pillars])

            if total_weight == 0:
                run.overall_inherent_score = 0.0
                run.overall_residual_score = 0.0
                run.overall_risk_level = 'low'
                continue

            # Calculate weighted average inherent score
            weighted_inherent_sum = sum([
                p.inherent_score * pillar_weights.get(p.pillar, 0)
                for p in core_pillars
            ])
            run.overall_inherent_score = weighted_inherent_sum / total_weight

            # Calculate weighted average residual score
            weighted_residual_sum = sum([
                p.residual_score * pillar_weights.get(p.pillar, 0)
                for p in core_pillars
            ])
            run.overall_residual_score = weighted_residual_sum / total_weight

            # Classify risk level based on residual score
            residual = run.overall_residual_score
            if residual < 1.7:
                run.overall_risk_level = 'low'
            elif residual < 2.4:
                run.overall_risk_level = 'medium'
            else:
                run.overall_risk_level = 'high'

    @api.constrains('period_start', 'period_end')
    def _check_period_dates(self):
        """Validate period dates."""
        for run in self:
            if run.period_start and run.period_end and run.period_start > run.period_end:
                raise ValidationError(_('Period start date must be before end date.'))

    def action_start_assessment(self):
        """Start the EWRA assessment."""
        # Auto-create pillars if none exist
        if not self.pillar_ids:
            self._create_default_pillars()

        self.write({'status': 'in_progress'})
        self.message_post(
            body=_('EWRA assessment started by %s') % self.env.user.name,
            subject=_('Assessment Started')
        )

    def _create_default_pillars(self):
        """Create default risk pillars for the assessment."""
        # Get default control band from settings or use 'adequate'
        default_band = 'adequate'
        if self.settings_snapshot_id:
            default_band = self.settings_snapshot_id.default_control_band

        # Get all customers for auto-population
        partners = self.env['res.partner'].search([
            ('customer_rank', '>', 0),
            ('is_company', '=', True)
        ])

        if not partners:
            # No customers, create empty pillars
            pillars_data = [
                {'pillar': 'customer', 'sequence': 10},
                {'pillar': 'geography', 'sequence': 20},
                {'pillar': 'products', 'sequence': 30},
                {'pillar': 'delivery', 'sequence': 40},
            ]
            for data in pillars_data:
                self.env['nexaml.ewra.pillar'].create({
                    'run_id': self.id,
                    'pillar': data['pillar'],
                    'sequence': data['sequence'],
                    'low_pct': 60.0,
                    'medium_pct': 30.0,
                    'high_pct': 10.0,
                    'control_band': default_band,
                })
        else:
            # Auto-populate from customer data
            total = len(partners)

            # Customer Risk
            customer_low = len(partners.filtered(lambda p: p.risk_level == 'low'))
            customer_medium = len(partners.filtered(lambda p: p.risk_level == 'medium'))
            customer_high = len(partners.filtered(lambda p: p.risk_level == 'high'))

            # Geography Risk
            geo_low = len(partners.filtered(lambda p: p.geography_risk < 1.7))
            geo_medium = len(partners.filtered(lambda p: 1.7 <= p.geography_risk < 2.4))
            geo_high = len(partners.filtered(lambda p: p.geography_risk >= 2.4))

            # Products Risk
            prod_low = len(partners.filtered(lambda p: p.product_risk < 1.7))
            prod_medium = len(partners.filtered(lambda p: 1.7 <= p.product_risk < 2.4))
            prod_high = len(partners.filtered(lambda p: p.product_risk >= 2.4))

            # Delivery Channel Risk
            del_low = len(partners.filtered(lambda p: p.channel_risk < 1.7))
            del_medium = len(partners.filtered(lambda p: 1.7 <= p.channel_risk < 2.4))
            del_high = len(partners.filtered(lambda p: p.channel_risk >= 2.4))

            pillars = [
                {'pillar': 'customer', 'sequence': 10, 'low': customer_low, 'medium': customer_medium, 'high': customer_high},
                {'pillar': 'geography', 'sequence': 20, 'low': geo_low, 'medium': geo_medium, 'high': geo_high},
                {'pillar': 'products', 'sequence': 30, 'low': prod_low, 'medium': prod_medium, 'high': prod_high},
                {'pillar': 'delivery', 'sequence': 40, 'low': del_low, 'medium': del_medium, 'high': del_high},
            ]

            for data in pillars:
                self.env['nexaml.ewra.pillar'].create({
                    'run_id': self.id,
                    'pillar': data['pillar'],
                    'sequence': data['sequence'],
                    'low_pct': round(data['low'] / total * 100, 1) if total > 0 else 33.3,
                    'medium_pct': round(data['medium'] / total * 100, 1) if total > 0 else 33.3,
                    'high_pct': round(data['high'] / total * 100, 1) if total > 0 else 33.4,
                    'control_band': default_band,
                })

        _logger.info('Auto-created %d risk pillars for EWRA run %s', len(self.pillar_ids), self.name)

    def action_complete(self):
        """Complete the EWRA assessment."""
        # Validate all pillars are assessed
        if not self.pillar_ids:
            raise ValidationError(_('Cannot complete EWRA without risk pillars.'))

        self.write({
            'status': 'completed',
            'approved_by': self.env.user.id,
            'approved_date': fields.Datetime.now()
        })
        self.message_post(
            body=_('EWRA assessment completed by %s') % self.env.user.name,
            subject=_('Assessment Completed')
        )

    def action_archive(self):
        """Archive the EWRA run."""
        self.write({'status': 'archived'})

    def action_initialize_pillars(self):
        """Manually initialize risk pillars."""
        if self.pillar_ids:
            raise ValidationError(_('Pillars already exist. Delete existing pillars first if you want to re-initialize.'))
        self._create_default_pillars()
        return {
            'type': 'ir.actions.client',
            'tag': 'display_notification',
            'params': {
                'title': _('Pillars Initialized'),
                'message': _('Risk pillars have been auto-populated from customer data.'),
                'type': 'success',
                'sticky': False,
            }
        }

    def action_generate_pdf(self):
        """Generate EWRA PDF report."""
        return self.env.ref('nexaml.action_report_ewra').report_action(self)

    def action_generate_docx(self):
        """Generate comprehensive EWRA DOCX report matching PDF."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement
            from docx.oxml.ns import qn
            from io import BytesIO
            import base64
        except ImportError:
            from odoo.exceptions import UserError
            raise UserError(_('python-docx library not installed. Please install with: pip install python-docx'))

        def clean_html(text):
            if not text:
                return ''
            return re.sub('<[^<]+?>', '', text)

        def set_cell_bg(cell, color):
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), color)
            cell._element.get_or_add_tcPr().append(shading)

        # Initialize document
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)

        # Get data
        company = self.company_id
        pillars = self.pillar_ids.sorted(lambda p: ['customer', 'geography', 'products', 'delivery', 'supplier'].index(p.pillar) if p.pillar in ['customer', 'geography', 'products', 'delivery', 'supplier'] else 99)
        customer_pillar = pillars.filtered(lambda p: p.pillar == 'customer')[:1]
        geography_pillar = pillars.filtered(lambda p: p.pillar == 'geography')[:1]
        products_pillar = pillars.filtered(lambda p: p.pillar == 'products')[:1]
        delivery_pillar = pillars.filtered(lambda p: p.pillar == 'delivery')[:1]
        supplier_pillar = pillars.filtered(lambda p: p.pillar == 'supplier')[:1]
        include_supplier = bool(supplier_pillar)
        settings = self.settings_snapshot_id
        pillar_weights = {'customer': 0.4, 'geography': 0.2, 'products': 0.25, 'delivery': 0.15, 'supplier': 0}
        core_pillars = pillars.filtered(lambda p: p.pillar != 'supplier')
        total_weight = sum([pillar_weights.get(p.pillar, 0) for p in core_pillars])
        weighted_sum = sum([p.residual_score * pillar_weights.get(p.pillar, 0) for p in core_pillars])
        avg_residual = weighted_sum / total_weight if total_weight > 0 else 0
        medium_threshold = settings.risk_threshold_medium if settings else 1.7
        high_threshold = settings.risk_threshold_high if settings else 2.4
        overall_label = 'High' if avg_residual >= high_threshold else ('Medium' if avg_residual >= medium_threshold else 'Low')

        # COVER PAGE
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('COMPLIANCE REPORT')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(71, 85, 105)
        run.font.bold = True

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Enterprise-Wide Risk Assessment')
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 41, 59)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Money Laundering / Terrorism Financing /\nProliferation Financing Risk Assessment (EWRA)')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(51, 65, 85)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'\n{self.period_start.strftime("%B %Y")}')
        run.font.size = Pt(12)
        run.font.bold = True

        doc.add_paragraph('\n\n')
        doc.add_paragraph(f'Company Name: {company.name}').runs[0].font.bold = True
        doc.add_paragraph(f'Assessment Period: {self.period_start.strftime("%d %b %Y")} to {self.period_end.strftime("%d %b %Y")}')
        doc.add_paragraph('\n')
        doc.add_paragraph('Confidential & Proprietary')
        p = doc.add_paragraph('Powered by nex.systems')
        p.runs[0].font.color.rgb = RGBColor(59, 130, 246)
        p.runs[0].font.bold = True
        doc.add_page_break()

        # TABLE OF CONTENTS
        heading = doc.add_heading('CONTENTS', 1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_items = [
            '1. REPORT INFORMATION', '2. FOREWORD BY THE COMPLIANCE OFFICER', '3. COMPANY OVERVIEW',
            '4. PURPOSE OF RISK ASSESSMENT', '5. FREQUENCY OF RISK ASSESSMENT', '6. RISK ASSESSMENT PROCESS',
            '7. RISK FACTORS', '8. RISK VISUALIZATION', '9. CUSTOMER RISK',
        ]
        if include_supplier:
            toc_items.append('10. SUPPLIER RISK')
        toc_items.extend([
            f'{11 if include_supplier else 10}. GEOGRAPHICAL RISK',
            f'{12 if include_supplier else 11}. PRODUCT RISK',
            f'{13 if include_supplier else 12}. TRANSACTION & PAYMENT MODE RISK',
            f'{14 if include_supplier else 13}. ONBOARDING CHANNEL RISK',
            f'{15 if include_supplier else 14}. INHERENT RISK ASSESSMENT RESULT',
            f'{16 if include_supplier else 15}. INTERNAL CONTROLS',
            f'{17 if include_supplier else 16}. SETTINGS SNAPSHOT (APPENDIX)',
            f'{18 if include_supplier else 17}. CONCLUSION',
        ])
        for item in toc_items:
            doc.add_paragraph(item, style='List Bullet')
        doc.add_page_break()

        # 1. REPORT INFORMATION
        heading = doc.add_heading('1. REPORT INFORMATION', 1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_table = doc.add_table(rows=8, cols=2)
        info_table.style = 'Light Grid Accent 1'

        # Report Title
        info_table.rows[0].cells[0].merge(info_table.rows[0].cells[1])
        cell = info_table.rows[0].cells[0]
        cell.text = 'Report Title'
        set_cell_bg(cell, 'CBD5E1')
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        info_table.rows[1].cells[0].merge(info_table.rows[1].cells[1])
        cell = info_table.rows[1].cells[0]
        cell.text = 'Enterprise-Wide Money Laundering / Terrorism Financing / Proliferation (ML/TF/PF) Risk Assessment Report'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Report Submitted Date
        info_table.rows[2].cells[0].merge(info_table.rows[2].cells[1])
        cell = info_table.rows[2].cells[0]
        cell.text = 'Report Submitted on'
        set_cell_bg(cell, 'CBD5E1')
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        info_table.rows[3].cells[0].merge(info_table.rows[3].cells[1])
        cell = info_table.rows[3].cells[0]
        cell.text = fields.Date.today().strftime('%d %B %Y')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Prepared By
        info_table.rows[4].cells[0].merge(info_table.rows[4].cells[1])
        cell = info_table.rows[4].cells[0]
        cell.text = 'ML/TF/PF Risk Assessment Report Prepared by'
        set_cell_bg(cell, 'CBD5E1')
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        info_table.rows[5].cells[0].text = 'Name:\n\nDesignation: Compliance Officer'
        info_table.rows[5].cells[1].text = 'Sign:\n\n\n'

        # Reviewed By
        info_table.rows[6].cells[0].merge(info_table.rows[6].cells[1])
        cell = info_table.rows[6].cells[0]
        cell.text = 'ML/TF/PF Risk Assessment Report Reviewed and Approved by'
        set_cell_bg(cell, 'CBD5E1')
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        info_table.rows[7].cells[0].text = 'Name:\nDesignation: Owner'
        info_table.rows[7].cells[1].text = 'Sign:\n\n\n'

        p = doc.add_paragraph('\n\nStamp:')
        p.runs[0].font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        # 2. FOREWORD BY THE COMPLIANCE OFFICER
        heading = doc.add_heading('2. FOREWORD BY THE COMPLIANCE OFFICER', 1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if self.narrative_id and self.narrative_id.foreword:
            doc.add_paragraph(clean_html(self.narrative_id.foreword))
        else:
            user_name = self.env.user.name
            p1 = doc.add_paragraph(f'I, {user_name}, the Compliance Officer of {company.name}, am honored to present this Enterprise-Wide Risk Assessment Report. This document provides a detailed overview of our company\'s ongoing efforts to identify, evaluate, and mitigate risks across all business operations. It reflects our steadfast commitment to excellence in governance, risk management, and regulatory compliance.')
            p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            p2 = doc.add_paragraph(f'Our risk assessment framework is designed to proactively detect potential threats and ensure effective controls and mitigation strategies are in place. {company.name} operates with strict adherence to applicable laws and regulations, including Anti-Money Laundering (AML), Counter Financing of Terrorism (CFT), and other relevant compliance frameworks. Through rigorous due diligence, continuous monitoring, and improvement, we uphold the integrity, transparency, and reliability of our operations.')
            p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            p3 = doc.add_paragraph(f'This report underscores our unwavering dedication to protecting our stakeholders\' interests and maintaining {company.name}\'s reputation as a responsible and compliant organization. By fostering a culture of accountability and compliance throughout the company, we ensure that risk management remains a cornerstone of our corporate governance.')
            p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            doc.add_paragraph('\nSincerely,')
            doc.add_paragraph(f'\n\n{user_name}\nCompliance Officer\n{company.name}').runs[0].font.bold = True
        doc.add_page_break()

        # 3. COMPANY OVERVIEW
        heading = doc.add_heading('3. COMPANY OVERVIEW', 1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        address_parts = [company.street, company.street2, company.city, company.state_id.name if company.state_id else '', company.zip]
        address = ', '.join(filter(None, address_parts))

        p1 = doc.add_paragraph(f'{company.name}, headquartered in {company.country_id.name or "its operating jurisdiction"}, operates in its designated industry and is licensed under {company.company_registry or "the relevant regulatory authorities"}. The company conducts its activities in accordance with applicable regulatory standards, maintaining full compliance with legal and operational requirements.')
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p2 = doc.add_paragraph(f'Operating primarily in the company\'s core sector, {company.name} upholds strong governance principles and transparent business practices. The organization\'s leadership structure fosters accountability, ethical conduct, and effective decision-making, ensuring the integrity and reliability of its operations.')
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p3 = doc.add_paragraph(f'With its headquarters located at {address}, and official contact channels via {company.email or "N/A"} and {company.phone or "N/A"}, the company serves both local and international clients. It uses {company.currency_id.name} as its default reporting currency, and adheres to all relevant tax obligations under registration number {company.vat or "N/A"}.')
        p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p4 = doc.add_paragraph(f'{company.name} remains dedicated to sustainable growth, compliance, and operational excellence - continuing to build trust and confidence among its stakeholders through responsible business practices.')
        p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_page_break()

        # 4. PURPOSE OF RISK ASSESSMENT
        heading = doc.add_heading('4. PURPOSE OF RISK ASSESSMENT', 1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p1 = doc.add_paragraph('To mitigate the risks associated with money laundering, terrorism financing, and proliferation financing (ML/TF/PF), a comprehensive risk-based assessment framework has been established. This framework facilitates the periodic and systematic evaluation of ML/TF/PF exposures, enabling the precise identification of risk distribution and the optimization of control mechanisms.')
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_paragraph('The key objectives and purposes of an EWRA:')
        doc.add_paragraph('Identification of Risks: The report serves as a tool to identify and document a broad range of risks that could impact the Company.', style='List Bullet')
        doc.add_paragraph('Prioritization of Risks: This prioritization helps the Company to focus their resources on managing the most significant and relevant risks.', style='List Bullet')
        doc.add_paragraph('Compliance and Regulatory Alignment: Compliance risks are thoroughly assessed and aligned with applicable regulatory requirements.', style='List Bullet')
        doc.add_paragraph('Strategic Decision Support: Comprehensive view of potential risks and it will help to the senior management and decision-makers.', style='List Bullet')
        doc.add_paragraph('Continuous Improvement: Regular updates to the EWRA allow the Company to adapt their risk management strategies to new challenges, emerging risks, or changes in the operating environment.', style='List Bullet')
        doc.add_paragraph('Demonstrating Due Diligence: This EWRA report will serve as evidence of the due diligence in identifying and addressing potential risks. This is particularly important for maintaining trust and credibility.', style='List Bullet')

        # 5. FREQUENCY OF RISK ASSESSMENT
        heading = doc.add_heading('5. FREQUENCY OF RISK ASSESSMENT', 1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p1 = doc.add_paragraph(f'The Company conducts a comprehensive ML/TF/PF risk assessment at least once annually, or whenever there are changes in business activities or updates to regulatory guidelines. This report presents the findings of the Company\'s risk assessment conducted for the reporting period of {self.period_start.strftime("%d %b %Y")} to {self.period_end.strftime("%d %b %Y")}.')
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_page_break()

        # 6. RISK ASSESSMENT PROCESS
        heading = doc.add_heading('6. RISK ASSESSMENT PROCESS', 1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph('The risk assessment process was conducted through the following steps:')
        doc.add_paragraph('Identification of Inherent Risks;', style='List Bullet')
        doc.add_paragraph('Evaluation of Risk Control and Mitigation Measures; and', style='List Bullet')
        doc.add_paragraph('Assessment of Residual Risks.', style='List Bullet')
        doc.add_page_break()

        # 7. RISK FACTORS
        heading = doc.add_heading('7. RISK FACTORS', 1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p1 = doc.add_paragraph('In adherence to regulatory requirements, the Company conduct a structured methodology to identify and classify its risk factors. These classifications are based on an assessment of both the probability and potential consequence of the risk. The specific risk factors evaluated include:')
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_paragraph('Customer Risk', style='List Bullet')
        if include_supplier:
            doc.add_paragraph('Supplier Risk', style='List Bullet')
        doc.add_paragraph('Geographical Risk', style='List Bullet')
        doc.add_paragraph('Products / Services Risk', style='List Bullet')
        doc.add_paragraph('Onboarding Channel Risk', style='List Bullet')

        # 8. RISK VISUALIZATION
        heading = doc.add_heading('8. RISK VISUALIZATION', 1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p1 = doc.add_paragraph('This section presents visual representations of the risk assessment findings to provide a clear and comprehensive overview of the organization\'s risk profile.')
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_heading('Overall Residual Risk', 3)
        p = doc.add_paragraph(f'{avg_residual:.2f} - {overall_label.upper()} RISK')
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(14)
        if overall_label == 'High':
            p.runs[0].font.color.rgb = RGBColor(153, 27, 27)
        elif overall_label == 'Medium':
            p.runs[0].font.color.rgb = RGBColor(133, 77, 14)
        else:
            p.runs[0].font.color.rgb = RGBColor(22, 101, 52)

        doc.add_heading('Risk Scores by Category', 3)
        viz_table = doc.add_table(rows=len(core_pillars) + 1, cols=5)
        viz_table.style = 'Light Grid Accent 1'

        # Header row
        viz_table.rows[0].cells[0].text = 'Risk Category'
        viz_table.rows[0].cells[1].text = 'Inherent'
        viz_table.rows[0].cells[2].text = 'Control %'
        viz_table.rows[0].cells[3].text = 'Residual'
        viz_table.rows[0].cells[4].text = 'Risk Level'

        for cell in viz_table.rows[0].cells:
            set_cell_bg(cell, 'CBD5E1')
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.bold = True

        for idx, pillar in enumerate(core_pillars, start=1):
            viz_table.rows[idx].cells[0].text = dict(pillar._fields['pillar'].selection).get(pillar.pillar, '')
            viz_table.rows[idx].cells[1].text = f'{pillar.inherent_score:.2f}'
            viz_table.rows[idx].cells[2].text = f'{pillar.control_pct:.0f}%'
            viz_table.rows[idx].cells[3].text = f'{pillar.residual_score:.2f}'
            level = dict(pillar._fields['residual_label'].selection).get(pillar.residual_label, 'N/A').upper()
            viz_table.rows[idx].cells[4].text = level
        doc.add_page_break()

        # 9. CUSTOMER RISK
        heading = doc.add_heading('9. CUSTOMER RISK', 1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if customer_pillar:
            p1 = doc.add_paragraph('The Company categorizes its customers into two main groups: Natural Persons (Individual Customers) and Legal Entities (Corporate Customers). Both domestic and international clients are engaged in accordance with international Anti-Money Laundering (AML) and Counter Financing of Terrorism (CFT) standards to ensure full regulatory compliance.')
            p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            doc.add_heading('Risk Distribution', 3)
            cust_table = doc.add_table(rows=4, cols=3)
            cust_table.style = 'Light Grid Accent 1'

            cust_table.rows[0].cells[0].text = 'Risk Level'
            cust_table.rows[0].cells[1].text = 'Percentage'
            cust_table.rows[0].cells[2].text = 'Note'
            for cell in cust_table.rows[0].cells:
                set_cell_bg(cell, 'CBD5E1')

            cust_table.rows[1].cells[0].text = 'Low'
            cust_table.rows[1].cells[1].text = f'{customer_pillar.low_pct:.1f}%'
            cust_table.rows[1].cells[2].text = 'Majority of customers fall into this category.'

            cust_table.rows[2].cells[0].text = 'Medium'
            cust_table.rows[2].cells[1].text = f'{customer_pillar.medium_pct:.1f}%'
            cust_table.rows[2].cells[2].text = 'Customers with some higher-risk attributes.'

            cust_table.rows[3].cells[0].text = 'High'
            cust_table.rows[3].cells[1].text = f'{customer_pillar.high_pct:.1f}%'
            cust_table.rows[3].cells[2].text = 'Customers requiring enhanced due diligence.'

            doc.add_heading('Customer Risk Analysis', 3)
            p2 = doc.add_paragraph(f'{company.name} performs thorough due diligence on all customers to effectively identify and manage potential risks. The process includes the following key measures:')
            p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            doc.add_paragraph('Customer Identity Verification: Confirming the authenticity of customers and validating Know Your Customer (KYC) documentation.', style='List Bullet')
            doc.add_paragraph('Background Screening: Conducting detailed checks on customer reputation, financial soundness, and legal standing, including adverse media, sanctions, and politically exposed person (PEP) screening.', style='List Bullet')
            doc.add_paragraph('Geographical Risk Assessment: Evaluating exposure linked to customers from high-risk or non-compliant jurisdictions.', style='List Bullet')
            doc.add_paragraph('Customer Profiling and Activity Review: Classifying customers by business type and nature to ensure consistency with regulatory and industry standards.', style='List Bullet')
            doc.add_paragraph('Regulatory Compliance Review: Ensuring all customers comply with AML/CFT obligations and maintain appropriate internal control systems.', style='List Bullet')
            doc.add_paragraph('Initial Risk Rating: Assigning preliminary risk levels based on jurisdiction, business category, transaction behavior, and compliance record.', style='List Bullet')
            doc.add_paragraph('Ongoing Monitoring: Continuously tracking customer transactions, behavioral changes, and regulatory compliance, followed by periodic risk reassessments.', style='List Bullet')

            customer_residual_label = 'High' if customer_pillar.residual_score >= high_threshold else ('Medium' if customer_pillar.residual_score >= medium_threshold else 'Low')
            p3 = doc.add_paragraph(f'Conclusion: Following the assessment, the Company\'s overall customer risk rating is determined as ')
            p3.add_run(customer_residual_label).bold = True
            p3.add_run('.')
            p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            doc.add_paragraph('Customer risk data not available.')
        doc.add_page_break()

        # 10. SUPPLIER RISK (conditional)
        if include_supplier:
            doc.add_heading('10. SUPPLIER RISK', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
            if supplier_pillar:
                doc.add_heading('Risk Distribution', 2)
                supp_table = doc.add_table(rows=4, cols=2)
                supp_table.style = 'Light Grid Accent 1'
                supp_table.rows[0].cells[0].text = 'Risk Level'
                supp_table.rows[0].cells[1].text = 'Percentage'
                supp_table.rows[1].cells[0].text = 'Low'
                supp_table.rows[1].cells[1].text = f'{supplier_pillar.low_pct:.1f}%'
                supp_table.rows[2].cells[0].text = 'Medium'
                supp_table.rows[2].cells[1].text = f'{supplier_pillar.medium_pct:.1f}%'
                supp_table.rows[3].cells[0].text = 'High'
                supp_table.rows[3].cells[1].text = f'{supplier_pillar.high_pct:.1f}%'
            else:
                doc.add_paragraph('Supplier risk data not available.')
            doc.add_page_break()

        # GEOGRAPHICAL RISK
        section_num = 11 if include_supplier else 10
        doc.add_heading(f'{section_num}. GEOGRAPHICAL RISK', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        if geography_pillar:
            doc.add_heading('Risk Distribution', 2)
            geo_table = doc.add_table(rows=4, cols=2)
            geo_table.style = 'Light Grid Accent 1'
            geo_table.rows[0].cells[0].text = 'Risk Level'
            geo_table.rows[0].cells[1].text = 'Percentage'
            geo_table.rows[1].cells[0].text = 'Low'
            geo_table.rows[1].cells[1].text = f'{geography_pillar.low_pct:.1f}%'
            geo_table.rows[2].cells[0].text = 'Medium'
            geo_table.rows[2].cells[1].text = f'{geography_pillar.medium_pct:.1f}%'
            geo_table.rows[3].cells[0].text = 'High'
            geo_table.rows[3].cells[1].text = f'{geography_pillar.high_pct:.1f}%'
        else:
            doc.add_paragraph('Geography risk data not available.')
        doc.add_page_break()

        # PRODUCT RISK
        section_num = 12 if include_supplier else 11
        doc.add_heading(f'{section_num}. PRODUCT RISK', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        if products_pillar:
            doc.add_heading('Risk Distribution', 2)
            prod_table = doc.add_table(rows=4, cols=2)
            prod_table.style = 'Light Grid Accent 1'
            prod_table.rows[0].cells[0].text = 'Risk Level'
            prod_table.rows[0].cells[1].text = 'Percentage'
            prod_table.rows[1].cells[0].text = 'Low'
            prod_table.rows[1].cells[1].text = f'{products_pillar.low_pct:.1f}%'
            prod_table.rows[2].cells[0].text = 'Medium'
            prod_table.rows[2].cells[1].text = f'{products_pillar.medium_pct:.1f}%'
            prod_table.rows[3].cells[0].text = 'High'
            prod_table.rows[3].cells[1].text = f'{products_pillar.high_pct:.1f}%'
        else:
            doc.add_paragraph('Product risk data not available.')
        doc.add_page_break()

        # TRANSACTION & PAYMENT MODE RISK
        section_num = 13 if include_supplier else 12
        doc.add_heading(f'{section_num}. TRANSACTION & PAYMENT MODE RISK', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        if delivery_pillar:
            doc.add_paragraph('Transaction and payment mode risks are assessed as part of delivery channel risk.')
        else:
            doc.add_paragraph('Data not available.')
        doc.add_page_break()

        # ONBOARDING CHANNEL RISK
        section_num = 14 if include_supplier else 13
        doc.add_heading(f'{section_num}. ONBOARDING CHANNEL RISK', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        if delivery_pillar:
            doc.add_heading('Risk Distribution', 2)
            del_table = doc.add_table(rows=4, cols=2)
            del_table.style = 'Light Grid Accent 1'
            del_table.rows[0].cells[0].text = 'Risk Level'
            del_table.rows[0].cells[1].text = 'Percentage'
            del_table.rows[1].cells[0].text = 'Low'
            del_table.rows[1].cells[1].text = f'{delivery_pillar.low_pct:.1f}%'
            del_table.rows[2].cells[0].text = 'Medium'
            del_table.rows[2].cells[1].text = f'{delivery_pillar.medium_pct:.1f}%'
            del_table.rows[3].cells[0].text = 'High'
            del_table.rows[3].cells[1].text = f'{delivery_pillar.high_pct:.1f}%'
        else:
            doc.add_paragraph('Delivery channel risk data not available.')
        doc.add_page_break()

        # INHERENT RISK ASSESSMENT RESULT
        section_num = 15 if include_supplier else 14
        doc.add_heading(f'{section_num}. INHERENT RISK ASSESSMENT RESULT', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f'Overall Inherent Risk Score: {self.overall_inherent_score:.2f}')
        doc.add_paragraph(f'Overall Risk Level: {overall_label.upper()}')
        doc.add_page_break()

        # INTERNAL CONTROLS
        section_num = 16 if include_supplier else 15
        heading = doc.add_heading(f'{section_num}. INTERNAL CONTROLS', 1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # (a) Customer and Supplier Onboarding Controls
        doc.add_heading('(a) Customer and Supplier Onboarding Controls', 3)
        p1 = doc.add_paragraph('The Company implemented due diligence and Enhanced Due Diligence (EDD) measures to verify customer and supplier identity, assess risk levels, and prevent financial crimes. These measures include:')
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_paragraph().add_run('Identity Verification:').bold = True
        doc.add_paragraph('Trade License, Certificate of Incorporation, and MOA/AOA.', style='List Bullet')
        doc.add_paragraph('For Ultimate Beneficial Owners (UBOs), shareholders, and authorized signatories\' passports, visas, and Emirates IDs.', style='List Bullet')
        doc.add_paragraph('Proof of address to confirm a legitimate business location.', style='List Bullet')
        doc.add_paragraph('VAT Certificate to confirm tax compliance.', style='List Bullet')
        doc.add_paragraph('Undertaking and PEP Declaration to assess corruption risks.', style='List Bullet')
        doc.add_paragraph('Completed KYC Form for comprehensive customer profiling.', style='List Bullet')

        # (b) Sanctions & Terrorist Financing Controls
        doc.add_heading('(b) Sanctions & Terrorist Financing Controls', 3)
        doc.add_paragraph().add_run('Screening Process:').bold = True
        doc.add_paragraph('Customers and suppliers are screened against global and UAE-specific sanction lists, including the UNSC, UAE Local Terrorist Lists, and other relevant lists.', style='List Bullet')
        doc.add_paragraph('PEP screening is conducted using the company\'s Nex Systems software to identify individuals with political exposure.', style='List Bullet')
        doc.add_paragraph('Continuous monitoring ensures sanctions lists are regularly updated.', style='List Bullet')

        # (c) Transaction Controls
        doc.add_heading('(c) Transaction Controls', 3)
        doc.add_paragraph('Only the approved profiles (customer and supplier) are eligible to execute the transactions.', style='List Bullet')
        doc.add_paragraph('EDD is conducted for all high-value transactions.', style='List Bullet')
        doc.add_paragraph('A system is in place to detect unusual or suspicious activities during transactions.', style='List Bullet')

        # (d) Ongoing Monitoring Controls
        doc.add_heading('(d) Ongoing Monitoring Controls', 3)
        doc.add_paragraph('The Company\'s screening software continuously monitors customer and supplier profiles. If there are any changes in risk status or matches with sanction or PEP categories, the system alerts the compliance officer.', style='List Bullet')
        doc.add_paragraph('In-depth verification and assessment for high-risk suppliers and customers, including those from high-risk jurisdictions.', style='List Bullet')
        doc.add_paragraph('Automated transaction monitoring tools detect unusual patterns, and escalation procedures are followed for high-risk profiles.', style='List Bullet')
        doc.add_paragraph('Regular updates and reviews of country risk classifications.', style='List Bullet')
        doc.add_paragraph('Adjustments to risk levels based on FATF and Know Your Country assessments.', style='List Bullet')
        doc.add_paragraph('If a customer\'s risk profile changes, a review is conducted, and high-risk customers undergo periodic re-screening.', style='List Bullet')
        doc.add_paragraph('Pattern analysis is conducted to detect unusual transaction behavior.', style='List Bullet')

        # (e) Authorized Cargo Service Provider Controls
        doc.add_heading('(e) Authorized Cargo Service Provider Controls', 3)
        doc.add_paragraph('The Company exclusively engages with verified and regulated cargo service providers to handle deliveries, ensuring secure and compliant movement of goods.', style='List Bullet')
        doc.add_paragraph('The delivery process is fully documented, and records are maintained for each transaction to ensure transparency and compliance.', style='List Bullet')
        doc.add_paragraph('No alternative delivery channels are permitted, reducing the risk of unauthorized transactions or illicit activities.', style='List Bullet')

        # (f) DPMS Reporting Controls
        doc.add_heading('(f) DPMS Reporting Controls', 3)
        doc.add_paragraph('Qualified transactions are reported on the GoAML portal as Dealers in Precious Metals and Stones Reports (DPMSR) within 14 days of the transaction execution date.', style='List Bullet')

        # (g) Reporting Mechanism Controls
        doc.add_heading('(g) Reporting Mechanism Controls', 3)
        doc.add_paragraph('Confirmed or partial matches during screening are reported to the UAE Financial Intelligence Unit (UAE-FIU) via the goAML system.', style='List Bullet')
        doc.add_paragraph('Suspicious transactions or suspicious activities are filed as Suspicious Transaction Reports (STRs) or Suspicious Activity Reports (SARs).', style='List Bullet')

        # (h) Restrictions Controls
        doc.add_heading('(h) Restrictions Controls', 3)
        doc.add_paragraph('Customers or suppliers cannot be onboarded without completing KYC and CDD procedures.', style='List Bullet')
        doc.add_paragraph('Transactions (buy/sell) cannot be initiated unless the required documents are provided by the customer or supplier.', style='List Bullet')
        doc.add_paragraph('Remote or third-party onboarding is not permitted.', style='List Bullet')

        # (i) System Based Controls
        doc.add_heading('(i) System Based Controls', 3)
        doc.add_paragraph('Also, the Company applies systematic risk mitigation process, the Compliance Officer can set the risk mitigation controls as per the requirements.', style='List Bullet')
        doc.add_paragraph('The assessment / risk mitigation will finalize the overall risk about the customer and the supplier based on the residual risk level the Compliance officer will set the ongoing monitoring for the high-risk profiles.', style='List Bullet')

        doc.add_heading('Control Effectiveness Summary', 2)
        ctrl_table = doc.add_table(rows=len(core_pillars) + 1, cols=3)
        ctrl_table.style = 'Light Grid Accent 1'
        ctrl_table.rows[0].cells[0].text = 'Risk Pillar'
        ctrl_table.rows[0].cells[1].text = 'Control Band'
        ctrl_table.rows[0].cells[2].text = 'Control %'

        for idx, pillar in enumerate(core_pillars, start=1):
            ctrl_table.rows[idx].cells[0].text = dict(pillar._fields['pillar'].selection).get(pillar.pillar, '')
            ctrl_table.rows[idx].cells[1].text = dict(pillar._fields['control_band'].selection).get(pillar.control_band, 'N/A')
            ctrl_table.rows[idx].cells[2].text = f'{pillar.control_pct:.0f}%'
        doc.add_page_break()

        # SETTINGS SNAPSHOT (APPENDIX)
        section_num = 17 if include_supplier else 16
        doc.add_heading(f'{section_num}. SETTINGS SNAPSHOT (APPENDIX)', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('This appendix captures the thresholds and control rules used for this EWRA run.')
        settings_table = doc.add_table(rows=6, cols=2)
        settings_table.style = 'Light Grid Accent 1'
        settings_table.rows[0].cells[0].text = 'Parameter'
        settings_table.rows[0].cells[1].text = 'Value'
        settings_table.rows[1].cells[0].text = 'Low Band'
        settings_table.rows[1].cells[1].text = '1.0 - 1.6'
        settings_table.rows[2].cells[0].text = 'Medium Band'
        settings_table.rows[2].cells[1].text = f'1.7 - {medium_threshold}'
        settings_table.rows[3].cells[0].text = 'High Band'
        settings_table.rows[3].cells[1].text = f'{high_threshold} - 3.0'
        settings_table.rows[4].cells[0].text = 'Controls Cap'
        settings_table.rows[4].cells[1].text = f'{settings.cap_pct if settings else 70}%'
        settings_table.rows[5].cells[0].text = 'Downgrade Threshold'
        settings_table.rows[5].cells[1].text = f'{settings.downgrade_threshold if settings else 35}%'
        doc.add_page_break()

        # CONCLUSION
        section_num = 18 if include_supplier else 17
        doc.add_heading(f'{section_num}. CONCLUSION', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        if self.narrative_id and self.narrative_id.conclusion:
            doc.add_paragraph(clean_html(self.narrative_id.conclusion))
        else:
            doc.add_paragraph(f'Based on the enterprise-wide risk assessment, the Company\'s overall residual ML/TF/PF risk level is assessed as {overall_label.upper()} with a risk score of {avg_residual:.2f}.')
            doc.add_paragraph('\nThe Company will continue to monitor and update the risk assessment as business conditions change.')

        if self.approved_by:
            doc.add_paragraph(f'\n\nApproved By: {self.approved_by.name}')
            doc.add_paragraph(f'Approval Date: {self.approved_date}')
            doc.add_paragraph('\n\nSignature: _______________________')

        # Save to BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        docx_data = output.read()
        output.close()

        # Create attachment
        filename = f'EWRA_{self.name}_{fields.Date.today()}.docx'
        attachment = self.env['ir.attachment'].create({
            'name': filename,
            'type': 'binary',
            'datas': base64.b64encode(docx_data),
            'res_model': 'nexaml.ewra.run',
            'res_id': self.id,
            'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        })

        return {
            'type': 'ir.actions.act_url',
            'url': f'/web/content/{attachment.id}?download=true',
            'target': 'self',
        }
