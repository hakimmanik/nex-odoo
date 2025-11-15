# -*- coding: utf-8 -*-
# Part of NexAML. See LICENSE file for full copyright and licensing details.

"""Comprehensive EWRA DOCX Report Generator."""

import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn


def clean_html(html_text):
    """Strip HTML tags from text."""
    if not html_text:
        return ''
    return re.sub('<[^<]+?>', '', html_text)


def set_cell_background(cell, color):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def generate_ewra_docx(ewra_run):
    """Generate comprehensive EWRA DOCX report.

    Args:
        ewra_run: nexaml.ewra.run record

    Returns:
        Document: python-docx Document object
    """
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # Get data
    company = ewra_run.company_id
    pillars = ewra_run.pillar_ids.sorted(lambda p: ['customer', 'geography', 'products', 'delivery', 'supplier'].index(p.pillar) if p.pillar in ['customer', 'geography', 'products', 'delivery', 'supplier'] else 99)
    customer_pillar = pillars.filtered(lambda p: p.pillar == 'customer')[:1]
    geography_pillar = pillars.filtered(lambda p: p.pillar == 'geography')[:1]
    products_pillar = pillars.filtered(lambda p: p.pillar == 'products')[:1]
    delivery_pillar = pillars.filtered(lambda p: p.pillar == 'delivery')[:1]
    supplier_pillar = pillars.filtered(lambda p: p.pillar == 'supplier')[:1]
    include_supplier = bool(supplier_pillar)
    settings = ewra_run.settings_snapshot_id

    # Calculate overall risk
    pillar_weights = {'customer': 0.4, 'geography': 0.2, 'products': 0.25, 'delivery': 0.15, 'supplier': 0}
    core_pillars = pillars.filtered(lambda p: p.pillar != 'supplier')
    total_weight = sum([pillar_weights.get(p.pillar, 0) for p in core_pillars])
    weighted_sum = sum([p.residual_score * pillar_weights.get(p.pillar, 0) for p in core_pillars])
    avg_residual = weighted_sum / total_weight if total_weight > 0 else 0
    medium_threshold = settings.risk_threshold_medium if settings else 1.7
    high_threshold = settings.risk_threshold_high if settings else 2.4
    overall_label = 'High' if avg_residual >= high_threshold else ('Medium' if avg_residual >= medium_threshold else 'Low')

    # ==================== COVER PAGE ====================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('COMPLIANCE REPORT')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(71, 85, 105)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Enterprise-Wide Risk Assessment')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 41, 59)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Money Laundering / Terrorism Financing /\nProliferation Financing Risk Assessment (EWRA)')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(51, 65, 85)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'\n{ewra_run.period_start.strftime("%B %Y")}')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(51, 65, 85)

    doc.add_paragraph('\n\n')
    p = doc.add_paragraph(f'Company Name: {company.name}')
    run = p.runs[0]
    run.font.size = Pt(10)
    run.font.bold = True

    p = doc.add_paragraph(f'Assessment Period: {ewra_run.period_start.strftime("%d %b %Y")} to {ewra_run.period_end.strftime("%d %b %Y")}')
    run = p.runs[0]
    run.font.size = Pt(9)

    doc.add_paragraph('\n\n')
    p = doc.add_paragraph('Confidential & Proprietary')
    run = p.runs[0]
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(51, 65, 85)

    p = doc.add_paragraph('Powered by nex.systems')
    run = p.runs[0]
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(59, 130, 246)
    run.font.bold = True

    doc.add_page_break()

    # ==================== TABLE OF CONTENTS ====================
    heading = doc.add_heading('CONTENTS', 1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    toc_items = [
        '1. REPORT INFORMATION',
        '2. FOREWORD BY THE COMPLIANCE OFFICER',
        '3. COMPANY OVERVIEW',
        '4. PURPOSE OF RISK ASSESSMENT',
        '5. FREQUENCY OF RISK ASSESSMENT',
        '6. RISK ASSESSMENT PROCESS',
        '7. RISK FACTORS',
        '8. RISK VISUALIZATION',
        '9. CUSTOMER RISK',
    ]
    if include_supplier:
        toc_items.append('10. SUPPLIER RISK')
    toc_items.extend([
        f'{11 if include_supplier else 10}. GEOGRAPHICAL RISK',
        f'{12 if include_supplier else 11}. PRODUCT RISK',
        f'{13 if include_supplier else 12}. TRANSACTION & PAYMENT MODE RISK',
        f'{14 if include_supplier else 13}. ONBOARDING CHANNEL RISK',
        f'{15 if include_supplier else 14}. INHERENT RISK ASSESSMENT RESULT',
        f'{16 if include_supplier else 15}. INTERNAL CONTROLS',
        f'{17 if include_supplier else 16}. SETTINGS SNAPSHOT (APPENDIX)',
        f'{18 if include_supplier else 17}. CONCLUSION',
    ])

    for item in toc_items:
        p = doc.add_paragraph(item)
        p.style = 'List Bullet'
        run = p.runs[0]
        run.font.size = Pt(10)

    doc.add_page_break()

    return doc, company, pillars, customer_pillar, geography_pillar, products_pillar, delivery_pillar, supplier_pillar, include_supplier, settings, avg_residual, overall_label, medium_threshold, high_threshold, core_pillars
